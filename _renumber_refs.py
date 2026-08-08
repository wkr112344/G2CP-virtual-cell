import docx, re

# 旧编号 -> 新编号（按正文首次出现顺序）
# 文档流: 引言[1-5] -> p19[6,7] -> 表3注[8] -> 3.5正文[9,10,11] -> 表4[12 Gao] -> 4.2[13 Lin]
MAP = {1:1, 10:2, 11:3, 12:4, 6:5, 2:6, 3:7, 5:8, 8:9, 13:10, 9:11, 7:12, 4:13}

# 新编号对应的旧编号顺序（参考文献重排）
NEW_ORDER_OLD = [1, 10, 11, 12, 6, 2, 3, 5, 8, 13, 9, 7, 4]

INTRO_EN = ("Predicting transcriptomic responses to molecular perturbations has emerged as a central challenge in "
            "AI-driven biology. UniPert-G2CP (Li et al., Cell, 2026) [1] unified genetic and chemical perturbation "
            "prediction in a single contrastive-embedding space across five cancer cell lines. Several approaches have "
            "been proposed for related tasks: scGen [2] uses variational autoencoders for single-cell perturbation "
            "response; CPA [3] uses an autoencoder framework for combinatorial perturbation prediction; ChemPert [4] "
            "predicts drug-induced transcriptomic changes via a multi-task model; and GEARS [5] introduced a "
            "graph-based framework integrating gene co-expression. To date, no independent reimplementation or "
            "large-scale extension of UniPert-G2CP has been reported.")
INTRO_CN = ("预测分子扰动后的转录组响应已成为 AI 驱动生物学中的核心挑战。UniPert-G2CP（Li 等，Cell，2026）[1] "
            "在单一对比嵌入空间中统一了五种癌细胞系上的遗传与化学扰动预测。针对相关任务已有多种方法被提出："
            "scGen [2] 使用变分自编码器预测单细胞扰动响应；CPA [3] 使用自编码器框架进行组合扰动预测；"
            "ChemPert [4] 通过多任务模型预测药物诱导的转录组变化；GEARS [5] 引入了整合基因共表达网络的图框架。"
            "迄今为止，尚未有对 UniPert-G2CP 的独立重新实现或大规模扩展的报道。")

def renumber_text(text):
    def repl(m):
        inner = m.group(1)
        nums = [int(x.strip()) for x in inner.split(',')]
        new = [str(MAP.get(n, n)) for n in nums]
        return '[' + ','.join(new) + ']'
    return re.sub(r'\[(\d+(?:\s*,\s*\d+)*)\]', repl, text)

def set_para_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)

def find_ref_span(paras):
    """定位参考文献段：以 [数字] 开头的连续段，返回 (start, count)"""
    starts = [i for i, p in enumerate(paras) if re.match(r'^\[\d+\]\s', p.text.strip())]
    if not starts:
        return None
    # 取最后一段连续块（参考文献在文末）
    blocks = []
    cur = [starts[0]]
    for s in starts[1:]:
        if s == cur[-1] + 1:
            cur.append(s)
        else:
            blocks.append(cur)
            cur = [s]
    blocks.append(cur)
    block = blocks[-1]  # 最后一块
    return block[0], len(block)

def process(path, intro_key, intro_new):
    d = docx.Document(path)
    paras = d.paragraphs

    span = find_ref_span(paras)
    if not span:
        print(f'!! {path}: 未找到参考文献段')
        return
    ref_start, ref_count = span
    print(f'  参考文献段: p{ref_start} 起 {ref_count} 段')

    # 1) 收集旧参考文献段落文本
    old_refs = [paras[ref_start + i].text for i in range(ref_count)]
    new_refs = []
    for new_idx, old_idx in enumerate(NEW_ORDER_OLD, start=1):
        txt = old_refs[old_idx - 1]
        txt = re.sub(r'^\[\d+\]', f'[{new_idx}]', txt.strip())
        new_refs.append(txt)

    # 2) 写回参考文献段落
    for i in range(ref_count):
        set_para_text(paras[ref_start + i], new_refs[i])

    # 3) 正文其余段落重编号（跳过参考文献）
    intro_idx = None
    for i, p in enumerate(paras):
        if ref_start <= i < ref_start + ref_count:
            continue
        t = p.text
        if intro_key in t:
            intro_idx = i
        if '[' in t:
            new_t = renumber_text(t)
            if new_t != t:
                set_para_text(p, new_t)

    # 4) 表格重编号
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    if '[' in p.text:
                        new_t = renumber_text(p.text)
                        if new_t != p.text:
                            set_para_text(p, new_t)

    # 5) 最后重写引言第一段（必须在步骤3之后）
    if intro_idx is not None:
        set_para_text(paras[intro_idx], intro_new)
        print(f'  引言重写 @p{intro_idx}')

    d.save(path)
    print(f'OK: {path}')

if __name__ == '__main__':
    process(r'C:\Users\wkr20\Desktop\biorxiv_english.docx',
            'Predicting transcriptomic responses', INTRO_EN)
    process(r'C:\Users\wkr20\Desktop\biorxiv_chinese.docx',
            '预测分子扰动后的转录组响应', INTRO_CN)
